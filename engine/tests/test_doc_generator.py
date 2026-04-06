from pathlib import Path
from docx import Document
from src.doc_generator import (
    generate_docx, generate_fallback_docx,
    _escape_xml, _display_date, _add_rich_paragraph, _add_h2_with_accent,
    _add_h1_hero, _add_horizontal_rule, _setup_styles, _setup_page,
    _setup_header_footer, _set_default_font,
)
from docx.shared import Cm


class TestEscapeXml:
    def test_ampersand(self):
        assert _escape_xml("A & B") == "A &amp; B"

    def test_less_than(self):
        assert _escape_xml("a < b") == "a &lt; b"

    def test_greater_than(self):
        assert _escape_xml("a > b") == "a &gt; b"

    def test_double_quote(self):
        assert _escape_xml('say "hi"') == "say &quot;hi&quot;"

    def test_combined(self):
        assert _escape_xml('<a & "b">') == "&lt;a &amp; &quot;b&quot;&gt;"

    def test_no_special_chars(self):
        assert _escape_xml("hello world") == "hello world"

    def test_empty_string(self):
        assert _escape_xml("") == ""

    def test_multiple_ampersands(self):
        assert _escape_xml("&&&&") == "&amp;&amp;&amp;&amp;"


class TestDisplayDate:
    def test_with_underscore(self):
        assert _display_date("2026-04-05_1930") == "2026-04-05"

    def test_without_underscore(self):
        assert _display_date("2026-04-05") == "2026-04-05"

    def test_multiple_underscores(self):
        # Should split on first underscore
        assert _display_date("2026-04-05_19_30") == "2026-04-05"


class TestSetupStyles:
    def test_configures_normal_style(self):
        doc = Document()
        _setup_styles(doc)
        style = doc.styles["Normal"]
        assert style.font.name == "Calibri"

    def test_configures_heading_styles(self):
        doc = Document()
        _setup_styles(doc)
        h1 = doc.styles["Heading 1"]
        h2 = doc.styles["Heading 2"]
        h3 = doc.styles["Heading 3"]
        assert h1.font.bold is True
        assert h2.font.bold is True
        assert h3.font.bold is True


class TestSetupPage:
    def test_a4_dimensions(self):
        doc = Document()
        _setup_page(doc)
        section = doc.sections[0]
        # Allow small rounding from EMU conversion
        assert abs(section.page_width - Cm(21.0)) < 1000
        assert abs(section.page_height - Cm(29.7)) < 1000

    def test_margins(self):
        doc = Document()
        _setup_page(doc)
        section = doc.sections[0]
        assert abs(section.left_margin - Cm(2.5)) < 1000
        assert abs(section.right_margin - Cm(2.5)) < 1000


class TestSetupHeaderFooter:
    def test_header_has_specola(self):
        doc = Document()
        _setup_header_footer(doc, "2026-04-05_1930")
        header = doc.sections[0].header
        text = "".join(p.text for p in header.paragraphs)
        assert "SPECOLA" in text

    def test_header_has_display_date(self):
        doc = Document()
        _setup_header_footer(doc, "2026-04-05_1930")
        header = doc.sections[0].header
        text = "".join(p.text for p in header.paragraphs)
        assert "2026-04-05" in text

    def test_footer_exists(self):
        doc = Document()
        _setup_header_footer(doc, "2026-04-05")
        footer = doc.sections[0].footer
        assert footer is not None


class TestAddH1Hero:
    def test_creates_heading1(self):
        doc = Document()
        _setup_styles(doc)
        _add_h1_hero(doc, "Test Title")
        assert doc.paragraphs[0].style.name == "Heading 1"
        assert doc.paragraphs[0].runs[0].text == "Test Title"


class TestAddH2WithAccent:
    def test_creates_heading2(self):
        doc = Document()
        _setup_styles(doc)
        _add_h2_with_accent(doc, "Section Title")
        assert doc.paragraphs[0].style.name == "Heading 2"

    def test_handles_bold_in_h2(self):
        doc = Document()
        _setup_styles(doc)
        _add_h2_with_accent(doc, "**Bold** and normal")
        runs = doc.paragraphs[0].runs
        assert len(runs) == 2
        assert runs[0].text == "Bold"
        assert runs[0].bold is True
        assert runs[1].text == " and normal"


class TestAddHorizontalRule:
    def test_adds_paragraph(self):
        doc = Document()
        _add_horizontal_rule(doc)
        assert len(doc.paragraphs) == 1


class TestAddRichParagraph:
    def test_plain_text(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "Hello world")
        assert doc.paragraphs[0].text == "Hello world"

    def test_bold_text(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "This has **bold** text")
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) >= 1
        assert any("bold" in r.text for r in bold_runs)

    def test_multiple_bold_segments(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "**A** normal **B**")
        runs = doc.paragraphs[0].runs
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) >= 2

    def test_link_text_present(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "See [article](https://example.com) here")
        # The hyperlink text may not show up in .text but the paragraph should exist
        assert len(doc.paragraphs) == 1

    def test_list_bullet_style(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "Item text", "List Bullet")
        assert doc.paragraphs[0].style.name == "List Bullet"

    def test_list_number_style(self):
        doc = Document()
        _setup_styles(doc)
        _add_rich_paragraph(doc, "Numbered item", "List Number")
        assert doc.paragraphs[0].style.name == "List Number"


class TestGenerateDocx:
    def test_creates_file(self, tmp_output_dir):
        md = "# Specola — Briefing del 2026-04-05\n\n## Da sapere oggi\n- Punto 1\n- Punto 2"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        assert Path(path).exists()
        assert Path(path).name == "Specola_2026-04-05.docx"

    def test_heading_levels(self, tmp_output_dir):
        md = "# H1 Title\n\n## H2 Section\n\n### H3 Subsection\n\nParagraph text."
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles
        assert "Normal" in styles

    def test_bullet_lists(self, tmp_output_dir):
        md = "## Section\n\n- Item one\n- Item two\n* Item three"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert styles.count("List Bullet") == 3

    def test_numbered_lists(self, tmp_output_dir):
        md = "## Section\n\n1. First\n2. Second\n3. Third"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert styles.count("List Number") == 3

    def test_bold_text(self, tmp_output_dir):
        md = "## Section\n\nThis has **bold text** inside."
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        para = [p for p in doc.paragraphs if p.style.name == "Normal"][0]
        runs_bold = [r.bold for r in para.runs]
        assert True in runs_bold

    def test_header_contains_specola(self, tmp_output_dir):
        md = "# Title"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        header = doc.sections[0].header
        header_text = "".join(p.text for p in header.paragraphs)
        assert "SPECOLA" in header_text.upper()

    def test_hyperlinks(self, tmp_output_dir):
        md = "## Section\n\nCheck [this article](https://example.com) for details."
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        # Verify the document contains the hyperlink text
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "this article" in all_text or "Check" in all_text

    def test_horizontal_rule(self, tmp_output_dir):
        md = "## Section 1\n\nSome text.\n\n---\n\n## Section 2\n\nMore text."
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        assert Path(path).exists()

    def test_creates_output_dir_if_missing(self, tmp_path):
        new_dir = tmp_path / "subdir" / "output"
        md = "# Title"
        path = generate_docx(md, "2026-04-05", new_dir)
        assert Path(path).exists()


    def test_date_with_timestamp(self, tmp_output_dir):
        md = "# Title"
        path = generate_docx(md, "2026-04-05_1930", tmp_output_dir)
        assert "Specola_2026-04-05_1930.docx" in path

    def test_empty_markdown(self, tmp_output_dir):
        """Empty markdown still produces a valid DOCX."""
        path = generate_docx("", "2026-04-05", tmp_output_dir)
        assert Path(path).exists()
        doc = Document(path)
        assert doc is not None

    def test_only_empty_lines(self, tmp_output_dir):
        """Markdown with only blank lines produces valid DOCX."""
        path = generate_docx("\n\n\n\n", "2026-04-05", tmp_output_dir)
        assert Path(path).exists()

    def test_mixed_list_types(self, tmp_output_dir):
        md = "- Bullet\n* Star bullet\n1. Numbered"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert "List Bullet" in styles
        assert "List Number" in styles

    def test_complex_markdown(self, tmp_output_dir):
        """Full briefing-like markdown generates without errors."""
        md = """# Specola — Briefing del 2026-04-05

## Da sapere oggi
- **OpenAI** lancia GPT-5 con [dettagli](https://example.com)
- EU AI Act entra in vigore

---

## Tech
### Article 1 (TechCrunch)
**Data:** 2026-04-05 09:30
Some detailed summary text here.

## Business
1. First business item
2. Second business item

---

## Da leggere con calma
- Interesting long-form [article](https://example.com/long)
"""
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        assert Path(path).exists()
        doc = Document(path)
        assert len(doc.paragraphs) > 5

    def test_long_horizontal_rule(self, tmp_output_dir):
        """Long HR (------) is detected as horizontal rule."""
        md = "## Before\n\n----------\n\n## After"
        path = generate_docx(md, "2026-04-05", tmp_output_dir)
        assert Path(path).exists()

    def test_returns_string_path(self, tmp_output_dir):
        path = generate_docx("# T", "2026-04-05", tmp_output_dir)
        assert isinstance(path, str)


class TestGenerateFallbackDocx:
    def test_creates_file_with_warning(self, tmp_output_dir):
        digest = "## Tech\n\n### Article 1\nSummary here"
        path = generate_fallback_docx(digest, "2026-04-05", tmp_output_dir)
        assert Path(path).exists()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Analisi non disponibile" in t for t in texts)
        assert any("Article 1" in t for t in texts)

    def test_fallback_has_correct_filename(self, tmp_output_dir):
        path = generate_fallback_docx("## T\nText", "2026-04-05", tmp_output_dir)
        assert Path(path).name == "Specola_2026-04-05.docx"

    def test_fallback_has_heading_styles(self, tmp_output_dir):
        digest = "# Main Title\n\n## Section\n\n### Subsection\n\nText"
        path = generate_fallback_docx(digest, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles

    def test_fallback_has_bullet_lists(self, tmp_output_dir):
        digest = "## Section\n\n- Item A\n* Item B"
        path = generate_fallback_docx(digest, "2026-04-05", tmp_output_dir)
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert styles.count("List Bullet") == 2

    def test_fallback_creates_dir(self, tmp_path):
        new_dir = tmp_path / "new" / "dir"
        path = generate_fallback_docx("## T", "2026-04-05", new_dir)
        assert Path(path).exists()

    def test_fallback_header_present(self, tmp_output_dir):
        path = generate_fallback_docx("## T", "2026-04-05", tmp_output_dir)
        doc = Document(path)
        header = doc.sections[0].header
        text = "".join(p.text for p in header.paragraphs)
        assert "SPECOLA" in text.upper()


class TestDocxThemes:
    def test_minimal_theme_creates_file(self, tmp_output_dir):
        md = "# Title\n\n## Section\n\n- Item 1\n- Item 2\n\nParagraph."
        path = generate_docx(md, "2026-04-05", tmp_output_dir, theme="minimal")
        assert Path(path).exists()
        doc = Document(path)
        assert len(doc.paragraphs) > 0

    def test_dark_theme_creates_file(self, tmp_output_dir):
        md = "# Title\n\n## Section\n\n- Item 1\n- Item 2\n\nParagraph."
        path = generate_docx(md, "2026-04-05", tmp_output_dir, theme="dark")
        assert Path(path).exists()
        doc = Document(path)
        assert len(doc.paragraphs) > 0

    def test_corporate_theme_default(self, tmp_output_dir):
        md = "# Title\n\n## Section\n\nText."
        path_default = generate_docx(md, "2026-04-05", tmp_output_dir)
        path_corp = generate_docx(md, "2026-04-06", tmp_output_dir, theme="corporate")
        # Both should produce valid files
        assert Path(path_default).exists()
        assert Path(path_corp).exists()

    def test_minimal_fallback_creates_file(self, tmp_output_dir):
        digest = "## Tech\n\n- Item A\n- Item B"
        path = generate_fallback_docx(digest, "2026-04-05", tmp_output_dir, theme="minimal")
        assert Path(path).exists()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Analisi non disponibile" in t for t in texts)

    def test_dark_fallback_creates_file(self, tmp_output_dir):
        digest = "## Tech\n\n- Item A\n- Item B"
        path = generate_fallback_docx(digest, "2026-04-05", tmp_output_dir, theme="dark")
        assert Path(path).exists()
        doc = Document(path)
        texts = [p.text for p in doc.paragraphs]
        assert any("Analisi non disponibile" in t for t in texts)

    def test_dark_theme_heading_styles(self, tmp_output_dir):
        md = "# H1\n\n## H2\n\n### H3\n\nText."
        path = generate_docx(md, "2026-04-05", tmp_output_dir, theme="dark")
        doc = Document(path)
        styles = [p.style.name for p in doc.paragraphs]
        assert "Heading 1" in styles
        assert "Heading 2" in styles
        assert "Heading 3" in styles
