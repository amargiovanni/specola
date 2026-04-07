"""DOCX generation from markdown — premium formatting."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_NUMBERED_RE = re.compile(r"^\d+\.\s+")
_HR_RE = re.compile(r"^-{3,}$")

# Color palette (defaults — kept for backward compatibility)
_NAVY = "1a1a2e"
_DARK_BLUE = "16213e"
_BLUE = "0f3460"
_ACCENT = "e94560"
_LIGHT_BG = "f0f3f8"
_LINK_COLOR = "2563eb"
_GRAY = "6b7280"
_LIGHT_GRAY = "d1d5db"

_DOCX_THEMES: dict[str, dict[str, str]] = {
    "corporate": {
        "navy": "1a1a2e",
        "dark_blue": "16213e",
        "blue": "0f3460",
        "accent": "e94560",
        "light_bg": "f0f3f8",
        "link_color": "2563eb",
        "gray": "6b7280",
        "light_gray": "d1d5db",
        "body_color": "374151",
        "warn_bg": "fef2f2",
        "warn_color": "dc2626",
    },
    "minimal": {
        "navy": "1a1a1a",
        "dark_blue": "1a1a1a",
        "blue": "1a1a1a",
        "accent": "a3a3a3",
        "light_bg": "fafaf9",
        "link_color": "1e3a5f",
        "gray": "737373",
        "light_gray": "d4d4d4",
        "body_color": "1a1a1a",
        "warn_bg": "fef2f2",
        "warn_color": "dc2626",
    },
    "dark": {
        "navy": "e0e0e0",
        "dark_blue": "c0c0d0",
        "blue": "a0a0b8",
        "accent": "e94560",
        "light_bg": "16213e",
        "link_color": "82b1ff",
        "gray": "8888aa",
        "light_gray": "3a3a5e",
        "body_color": "e0e0e0",
        "warn_bg": "2e1a1a",
        "warn_color": "ff6b6b",
    },
}


def _get_theme_colors(theme: str = "corporate") -> dict[str, str]:
    """Return color palette for the given theme."""
    return _DOCX_THEMES.get(theme, _DOCX_THEMES["corporate"])


def _set_default_font(doc: Document, colors: dict[str, str] | None = None) -> None:
    """Set document default font."""
    c = colors or _get_theme_colors()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor.from_string(c["body_color"])
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.3


def _configure_heading_style(
    doc: Document, level: int, size: int, color_hex: str,
    space_before: int, space_after: int, all_caps: bool = False,
) -> None:
    """Configure heading style."""
    style = doc.styles[f"Heading {level}"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.bold = True
    font.color.rgb = RGBColor.from_string(color_hex)
    font.all_caps = all_caps
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.2
    pf.keep_with_next = True


def _setup_styles(doc: Document, colors: dict[str, str] | None = None) -> None:
    """Configure all document styles."""
    c = colors or _get_theme_colors()
    _set_default_font(doc, colors=c)
    _configure_heading_style(doc, 1, 22, c["navy"], 0, 10)
    _configure_heading_style(doc, 2, 14, c["dark_blue"], 20, 6)
    _configure_heading_style(doc, 3, 11, c["blue"], 12, 4)

    # List Bullet style
    try:
        lb = doc.styles["List Bullet"]
        lb.font.name = "Calibri"
        lb.font.size = Pt(10.5)
        lb.font.color.rgb = RGBColor.from_string(c["body_color"])
        lb.paragraph_format.space_after = Pt(4)
        lb.paragraph_format.space_before = Pt(2)
        lb.paragraph_format.line_spacing = 1.3
    except KeyError:
        pass

    # List Number style
    try:
        ln = doc.styles["List Number"]
        ln.font.name = "Calibri"
        ln.font.size = Pt(10.5)
        ln.font.color.rgb = RGBColor.from_string(c["body_color"])
        ln.paragraph_format.space_after = Pt(4)
        ln.paragraph_format.line_spacing = 1.3
    except KeyError:
        pass


def _setup_page(doc: Document) -> None:
    """Configure A4 page with refined margins."""
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def _display_date(date: str) -> str:
    """Extract display date from timestamp. '2026-04-05_1930' -> '2026-04-05'."""
    return date.split("_")[0] if "_" in date else date


def _setup_header_footer(doc: Document, date: str, colors: dict[str, str] | None = None) -> None:
    """Add styled header and footer."""
    c = colors or _get_theme_colors()
    section = doc.sections[0]
    display = _display_date(date)

    # Header — thin line with Specola branding
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()

    run_icon = hp.add_run("\u25C6  ")  # diamond symbol
    run_icon.font.size = Pt(8)
    run_icon.font.color.rgb = RGBColor.from_string(c["accent"])

    run_left = hp.add_run("SPECOLA")
    run_left.font.size = Pt(8)
    run_left.font.color.rgb = RGBColor.from_string(c["gray"])
    run_left.font.all_caps = True
    run_left.font.bold = True

    hp.add_run("  \u2502  ")  # vertical bar separator
    hp.runs[-1].font.size = Pt(8)
    hp.runs[-1].font.color.rgb = RGBColor.from_string(c["light_gray"])

    run_right = hp.add_run(display)
    run_right.font.size = Pt(8)
    run_right.font.color.rgb = RGBColor.from_string(c["gray"])

    # Bottom border on header paragraph
    pPr = hp._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{c["light_gray"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top border on footer
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{c["light_gray"]}"/>'
        f'</w:pBdr>'
    )
    fpPr.append(fpBdr)

    run = fp.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(c["gray"])

    fld_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._r.append(fld_begin)
    instr = run._r.makeelement(qn("w:instrText"), {})
    instr.text = " PAGE "
    run._r.append(instr)
    fld_end = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run._r.append(fld_end)


def _add_hyperlink(paragraph, text: str, url: str, colors: dict[str, str] | None = None) -> None:
    """Add a clickable hyperlink to a paragraph."""
    c = colors or _get_theme_colors()
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # w:sz expects integer half-points (21 = 10.5pt)
    sz_half_pts = int(10.5 * 2)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:color w:val="{c["link_color"]}"/>'
        f'      <w:sz w:val="{sz_half_pts}"/>'
        f'      <w:szCs w:val="{sz_half_pts}"/>'
        f'      <w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/>'
        f'      <w:u w:val="single"/>'
        f'    </w:rPr>'
        f'    <w:t xml:space="preserve">{_escape_xml(text)}</w:t>'
        f'  </w:r>'
        f'</w:hyperlink>'
    )
    paragraph._p.append(hyperlink)


def _escape_xml(text: str) -> str:
    """Escape XML special characters."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")


def _add_rich_paragraph(doc: Document, text: str, style: str = "Normal", colors: dict[str, str] | None = None) -> None:
    """Add a paragraph with bold and hyperlink support."""
    c = colors or _get_theme_colors()
    para = doc.add_paragraph(style=style)

    # Split by links first, then handle bold within each segment
    segments = _LINK_RE.split(text)
    # segments: [before, link_text, link_url, between, link_text, link_url, after, ...]

    i = 0
    while i < len(segments):
        if i + 2 < len(segments) and i > 0 and (i % 3) == 1:
            # This is a link: segments[i] = text, segments[i+1] = url
            _add_hyperlink(para, segments[i], segments[i + 1], colors=c)
            i += 2
        else:
            # Regular text — handle bold
            chunk = segments[i]
            if chunk:
                bold_parts = _BOLD_RE.split(chunk)
                for j, part in enumerate(bold_parts):
                    if not part:
                        continue
                    run = para.add_run(part)
                    run.font.name = "Calibri"
                    run.font.size = Pt(10.5)
                    if j % 2 == 1:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(c["dark_blue"])
            i += 1


def _add_h2_with_accent(doc: Document, text: str, colors: dict[str, str] | None = None) -> None:
    """Add H2 heading with colored left border accent."""
    c = colors or _get_theme_colors()
    para = doc.add_paragraph(style="Heading 2")

    # Add bold parts handling
    bold_parts = _BOLD_RE.split(text)
    for j, part in enumerate(bold_parts):
        if not part:
            continue
        run = para.add_run(part)
        if j % 2 == 1:
            run.bold = True

    # Add left border accent
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{c["accent"]}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="{c["light_gray"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _add_h1_hero(doc: Document, text: str, colors: dict[str, str] | None = None) -> None:
    """Add H1 as a hero title with accent underline."""
    c = colors or _get_theme_colors()
    para = doc.add_paragraph(style="Heading 1")
    run = para.add_run(text)
    run.font.color.rgb = RGBColor.from_string(c["navy"])

    # Thick accent bottom border
    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="6" w:color="{c["accent"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # Add spacing after
    para.paragraph_format.space_after = Pt(16)


def _add_horizontal_rule(doc: Document, colors: dict[str, str] | None = None) -> None:
    """Add a subtle horizontal rule."""
    c = colors or _get_theme_colors()
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(12)

    pPr = para._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{c["light_gray"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def generate_docx(markdown: str, date: str, output_dir: str | Path, theme: str = "corporate") -> str:
    """Generate premium DOCX from markdown. Returns output file path."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"Specola_{date}.docx"

    colors = _get_theme_colors(theme)

    doc = Document()
    _setup_styles(doc, colors=colors)
    _setup_page(doc)
    _setup_header_footer(doc, date, colors=colors)

    for line in markdown.split("\n"):
        stripped = line.strip()

        if not stripped:
            continue

        # Horizontal rule
        if _HR_RE.match(stripped):
            _add_horizontal_rule(doc, colors=colors)
        # Headings
        elif stripped.startswith("### "):
            para = doc.add_paragraph(stripped[4:], "Heading 3")
        elif stripped.startswith("## "):
            _add_h2_with_accent(doc, stripped[3:], colors=colors)
        elif stripped.startswith("# "):
            _add_h1_hero(doc, stripped[2:], colors=colors)
        # Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_rich_paragraph(doc, stripped[2:], "List Bullet", colors=colors)
        elif _NUMBERED_RE.match(stripped):
            text = _NUMBERED_RE.sub("", stripped)
            _add_rich_paragraph(doc, text, "List Number", colors=colors)
        # Regular paragraph
        else:
            _add_rich_paragraph(doc, stripped, colors=colors)

    doc.save(str(output_path))
    return str(output_path)


def generate_fallback_docx(digest: str, date: str, output_dir: str | Path, theme: str = "corporate") -> str:
    """Generate fallback DOCX from raw digest with warning."""
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"Specola_{date}.docx"

    colors = _get_theme_colors(theme)

    doc = Document()
    _setup_styles(doc, colors=colors)
    _setup_page(doc)
    _setup_header_footer(doc, date, colors=colors)

    # Warning banner
    warn = doc.add_paragraph()
    warn.paragraph_format.space_after = Pt(16)
    pPr = warn._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["warn_bg"]}" w:val="clear"/>')
    pPr.append(shd)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="18" w:space="8" w:color="{colors["warn_color"]}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    run = warn.add_run("\u26a0  Analisi non disponibile. Di seguito il digest grezzo.")
    run.font.color.rgb = RGBColor.from_string(colors["warn_color"])
    run.font.bold = True
    run.font.size = Pt(11)

    for line in digest.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], "Heading 3")
        elif stripped.startswith("## "):
            _add_h2_with_accent(doc, stripped[3:], colors=colors)
        elif stripped.startswith("# "):
            _add_h1_hero(doc, stripped[2:], colors=colors)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            _add_rich_paragraph(doc, stripped[2:], "List Bullet", colors=colors)
        else:
            _add_rich_paragraph(doc, stripped, colors=colors)

    doc.save(str(output_path))
    return str(output_path)
